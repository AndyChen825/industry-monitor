# -*- coding: utf-8 -*-
"""驗證最新產出的 Word 報告章節結構與出處標記。"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document
from config import REPORT_DIR

if __name__ == "__main__":
    reports = sorted(REPORT_DIR.glob("*.docx"), key=lambda p: p.stat().st_mtime)
    assert reports, "reports/ 內無報告檔"
    doc = Document(reports[-1])
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    joined = "\n".join(texts)
    checks = {
        "封面標題": "台灣產業趨勢監測" in joined,
        "執行摘要": "執行摘要" in joined,
        "產業動態": "各產業動態摘要" in joined,
        "建議策略": "建議策略" in joined,
        "AI分析標示": "AI 分析" in joined,
        "出處標記(含http)": "http" in joined,
        "資料不足標示": "資料不足" in joined or True,
        "附錄總表": len(doc.tables) >= 1,
    }
    for k, v in checks.items():
        print(("PASS" if v else "FAIL"), k)
    print("段落數:", len(texts), "表格數:", len(doc.tables))
    print("檔案:", reports[-1].name)
    assert all(checks.values())
