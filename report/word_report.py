# -*- coding: utf-8 -*-
"""Word 報告產生器(python-docx)。

結構:封面 → 執行摘要與總體結論(附出處) → 統計總覽(產業篇數/公司聲量)
     → 各產業動態 → 建議策略(標明 AI 分析) → 資料缺口說明 → 附錄:引用來源總表。
"""
from collections import Counter
from datetime import datetime, timezone, timedelta

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import INDUSTRIES, REPORT_DIR
from analysis.summarizer import summarize_period, detect_themes, format_citation

SECTION_NUMERALS = ["壹", "貳", "參", "肆", "伍", "陸", "柒"]

TZ_TAIPEI = timezone(timedelta(hours=8))

MAIN_COLOR = RGBColor(0x1F, 0x4E, 0x79)   # 深藍標題
GRAY = RGBColor(0x66, 0x66, 0x66)


def _set_base_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    # 中文字型
    style.element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "微軟正黑體")


def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = MAIN_COLOR
    return h


def _cite_para(doc, article, bullet=True):
    """一則條目:標題 + 出處(媒體、日期、連結)。"""
    p = doc.add_paragraph(style="List Bullet" if bullet else None)
    p.add_run(article["title"])
    cite = p.add_run(" " + format_citation(article))
    cite.font.size = Pt(9)
    cite.font.color.rgb = GRAY
    return p


def build_report(articles, start_date, end_date, fetch_status=None,
                 title="台灣產業趨勢監測週報", company_section=None, out_path=None,
                 watchlist=None, articles_prev=None):
    """產出 Word 報告,回傳檔案路徑。

    articles:期間內文章(dict list,含 industry 欄位)
    fetch_status:各來源最近抓取狀態(標示資料缺口)
    company_section:公司情報 dict(公司報告用),含 name/registry/news
    watchlist:公司觀察清單(有提供時產出各產業公司聲量統計)
    articles_prev:等長前一期文章(有提供時 SOV 對比附前期變化)
    """
    doc = Document()
    _set_base_font(doc)
    now = datetime.now(TZ_TAIPEI)

    # ---- 統計基礎(供執行摘要/統計總覽共用) ----
    classified = [a for a in articles if a.get("industry")]
    ind_counts = Counter(a["industry"] for a in classified)
    comp_stats = {}
    if watchlist:
        from analysis.companies import company_counts
        comp_stats = company_counts(articles, watchlist, top_n=10)
    themes = detect_themes(articles)
    section_no = iter(SECTION_NUMERALS)

    # ---- 封面 ----
    for _ in range(6):
        doc.add_paragraph()
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run(title)
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = MAIN_COLOR
    for text in (f"報告期間:{start_date} ~ {end_date}",
                 f"產出日期:{now.strftime('%Y-%m-%d %H:%M')}(台北時間)",
                 "資料來源:公開新聞媒體 RSS、經濟部商工登記開放資料"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p.add_run(text)
        rr.font.size = Pt(12)
        rr.font.color.rgb = GRAY
    doc.add_page_break()

    # ---- 執行摘要與總體結論 ----
    _heading(doc, f"{next(section_no)}、執行摘要與總體結論", 1)
    by_ind = summarize_period(articles, top_n=3)
    if not by_ind:
        doc.add_paragraph("本期間內未取得可歸類之產業新聞資料 —— 資料不足。")
    else:
        # 總體結論(依統計自動生成,數據可溯源)
        _heading(doc, "一、總體結論(系統依統計自動生成)", 2)
        total_classified = sum(ind_counts.values())
        top_inds = ind_counts.most_common(3)
        parts = [
            f"本期間({start_date} ~ {end_date})共蒐集 {len(articles)} 篇公開報導,"
            f"其中 {total_classified} 篇可歸類至 {len(ind_counts)} 個產業類別"
            f"({len(articles) - total_classified} 篇為一般性新聞未歸類)。",
            "聲量最高的產業依序為:" + "、".join(
                f"{ind}({cnt} 篇,佔可歸類報導 {cnt / total_classified * 100:.1f}%)"
                for ind, cnt in top_inds) + "。",
        ]
        if comp_stats:
            merged = {}
            for ranked in comp_stats.values():
                for name, cnt in ranked:
                    merged[name] = max(merged.get(name, 0), cnt)
            top_comps = sorted(merged.items(), key=lambda x: -x[1])[:5]
            parts.append("公司層面,全市場被提及次數最高者為:" + "、".join(
                f"{n}({c} 篇)" for n, c in top_comps) + "。")
        if themes:
            parts.append("跨產業主題方面,以「" + themes[0][0] +
                         f"」相關報導最多({themes[0][1]} 篇)" +
                         ("、其次為「" + themes[1][0] + f"」({themes[1][1]} 篇)"
                          if len(themes) > 1 else "") + ",詳見建議策略章節。")
        doc.add_paragraph(" ".join(parts))

        _heading(doc, "二、各產業重點(每點附出處)", 2)
        for ind in INDUSTRIES:
            arts = by_ind.get(ind)
            if not arts:
                continue
            p = doc.add_paragraph()
            p.add_run(f"【{ind}】").bold = True
            for a in arts[:2]:
                _cite_para(doc, a)

    # ---- 統計總覽 ----
    doc.add_page_break()
    _heading(doc, f"{next(section_no)}、統計總覽", 1)
    _heading(doc, "一、各產業報導篇數", 2)
    if not ind_counts:
        doc.add_paragraph("本期間無可歸類報導 —— 資料不足。")
    else:
        total_classified = sum(ind_counts.values())
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, t in enumerate(("產業", "報導篇數", "佔可歸類報導比例")):
            hdr[i].text = t
        for ind in INDUSTRIES:
            cnt = ind_counts.get(ind, 0)
            row = table.add_row().cells
            row[0].text = ind
            row[1].text = str(cnt)
            row[2].text = f"{cnt / total_classified * 100:.1f}%" if cnt else "—"
    _heading(doc, "二、各產業公司聲量(前 10 大)", 2)
    if not comp_stats:
        doc.add_paragraph("未提供公司觀察清單,本節從缺。")
    else:
        doc.add_paragraph(
            "統計方式:以上市櫃公司名錄+自訂品牌清單比對報導之標題/摘要/來源,"
            "計算各公司被提及之篇數(不限文章之產業分類;同名詞彙可能造成少量誤計)。")
        for ind in INDUSTRIES:
            ranked = comp_stats.get(ind)
            if not ranked:
                continue
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{ind}:").bold = True
            p.add_run("、".join(f"{n}({c} 篇)" for n, c in ranked))

    # 產業篇數圖 + 每日趨勢圖(圖表失敗不影響報告產出)
    try:
        from report.charts import industry_bar_png, daily_trend_png
        if ind_counts:
            doc.add_picture(industry_bar_png(dict(ind_counts)), width=Cm(16))
        doc.add_picture(daily_trend_png(articles, start_date, end_date), width=Cm(16))
    except Exception:  # noqa: BLE001 — 無 matplotlib/字型等環境問題時略過圖表
        doc.add_paragraph("(圖表產生失敗,僅提供表格數據)")

    # 統計總覽 三、競品 SOV 對比(固定對比組)
    _heading(doc, "三、競品聲量對比(SOV)", 2)
    from config import REPORT_SOV_GROUPS
    from analysis.companies import make_matcher

    def _find_aliases(name):
        for comps in (watchlist or {}).values():
            if name in comps:
                return comps[name]
        return [name]

    def _mention_count(arts, match):
        return sum(1 for a in arts
                   if match(f"{a.get('title', '')} {a.get('summary', '')} {a.get('source', '')}"))

    doc.add_paragraph(
        "SOV(Share of Voice)= 該公司篇數 ÷ 對比組合計篇數。"
        + ("「變化」為相較等長前一期之篇數增減。" if articles_prev else ""))
    for group_name, names in REPORT_SOV_GROUPS.items():
        rows = []
        for n in names:
            m = make_matcher(_find_aliases(n))
            cur_c = _mention_count(articles, m)
            prev_c = _mention_count(articles_prev, m) if articles_prev else None
            rows.append((n, cur_c, prev_c))
        gtotal = sum(c for _, c, _ in rows)
        gp = doc.add_paragraph()
        gp.add_run(f"【{group_name}】合計 {gtotal} 篇").bold = True
        table = doc.add_table(rows=1, cols=4 if articles_prev else 3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        headers = ["公司", "本期篇數", "SOV"] + (["較前期變化"] if articles_prev else [])
        for i, t in enumerate(headers):
            hdr[i].text = t
        for n, c, pv in sorted(rows, key=lambda x: -x[1]):
            row = table.add_row().cells
            row[0].text = n
            row[1].text = str(c)
            row[2].text = f"{c / gtotal * 100:.1f}%" if gtotal else "—"
            if articles_prev:
                diff = c - (pv or 0)
                row[3].text = f"{'▲' if diff > 0 else ('▼' if diff < 0 else '—')}{abs(diff)}(前期 {pv})"

    # ---- 負面聲量警示 ----
    doc.add_page_break()
    _heading(doc, f"{next(section_no)}、負面聲量警示", 1)
    warn = doc.add_paragraph(
        "以下為規則式偵測(觀察名單公司 × 負面關鍵詞同時命中),僅供初步示警,"
        "文章實際立場需人工確認(例:報導同業出事亦可能命中)。")
    warn.runs[0].font.size = Pt(9)
    warn.runs[0].font.color.rgb = GRAY
    from analysis.alerts import negative_alerts, positive_signals
    alerts = negative_alerts(articles)
    if not alerts:
        doc.add_paragraph("本期未偵測到觀察名單公司之負面訊號。")
    else:
        for al in alerts[:10]:
            p = doc.add_paragraph()
            p.add_run(f"⚠ {al['company']}(疑似負面 {al['count']} 篇;"
                      f"命中詞:{'、'.join(al['keywords'])})").bold = True
            for s in al["samples"][:3]:
                sp = doc.add_paragraph(style="List Bullet")
                sp.add_run(s["title"])
                c = sp.add_run(f" ({s['source']},{s['published_at'] or '日期不明'},{s['url']})")
                c.font.size = Pt(9)
                c.font.color.rgb = GRAY

    # ---- 商機訊號(業務拜訪線索) ----
    _heading(doc, f"{next(section_no)}、商機訊號(業務拜訪線索)", 1)
    warn2 = doc.add_paragraph(
        "規則式偵測:觀察名單公司 × 正面事件詞(得標/擴廠/展店/導入AI 等)同時命中,"
        "代表該公司可能有預算與需求;僅供初步線索,請人工確認。")
    warn2.runs[0].font.size = Pt(9)
    warn2.runs[0].font.color.rgb = GRAY
    opps = positive_signals(articles)
    if not opps:
        doc.add_paragraph("本期未偵測到觀察名單公司之商機訊號。")
    else:
        for op in opps[:12]:
            p = doc.add_paragraph()
            p.add_run(f"💡 {op['company']}(商機訊號 {op['count']} 篇;"
                      f"命中詞:{'、'.join(op['keywords'])})").bold = True
            for s in op["samples"][:2]:
                sp = doc.add_paragraph(style="List Bullet")
                sp.add_run(s["title"])
                c = sp.add_run(f" ({s['source']},{s['published_at'] or '日期不明'},{s['url']})")
                c.font.size = Pt(9)
                c.font.color.rgb = GRAY

    # ---- 公司情報(公司報告專用) ----
    if company_section:
        doc.add_page_break()
        _heading(doc, f"{next(section_no)}、公司情報:{company_section['name']}", 1)
        reg = company_section.get("registry") or {}
        _heading(doc, "一、商工登記資料", 2)
        if reg.get("results"):
            table = doc.add_table(rows=0, cols=2)
            table.style = "Light Grid Accent 1"
            for comp in reg["results"][:3]:
                for k, v in comp.items():
                    row = table.add_row()
                    row.cells[0].text = k
                    row.cells[1].text = str(v)
                table.add_row()
        else:
            doc.add_paragraph("未能自動取得商工登記資料 —— 資料不足。")
        note = reg.get("note")
        if note:
            np_ = doc.add_paragraph(note)
            np_.runs[0].font.size = Pt(9)
            np_.runs[0].font.color.rgb = GRAY
        _heading(doc, "二、相關報導與公告", 2)
        news = company_section.get("news") or []
        if news:
            for a in news[:20]:
                _cite_para(doc, a)
        else:
            doc.add_paragraph("期間內未蒐集到該公司相關報導 —— 資料不足。")
        site_note = company_section.get("site_note")
        if site_note:
            sp = doc.add_paragraph("官網公告限制:" + site_note)
            sp.runs[0].font.size = Pt(9)
            sp.runs[0].font.color.rgb = GRAY

    # ---- 各產業動態 ----
    doc.add_page_break()
    _heading(doc, f"{next(section_no)}、各產業動態摘要", 1)
    by_ind_full = summarize_period(articles, top_n=8)
    for ind, spec in INDUSTRIES.items():
        _heading(doc, ind, 2)
        dg = doc.add_paragraph(f"主計總處行業統計分類對照:{spec['dgbas']}")
        dg.runs[0].font.size = Pt(9)
        dg.runs[0].font.color.rgb = GRAY
        arts = by_ind_full.get(ind)
        if not arts:
            doc.add_paragraph("本期間無可歸類至此產業之報導 —— 資料不足。")
            continue
        stat = doc.add_paragraph()
        stat.add_run(f"本期共 {ind_counts.get(ind, 0)} 篇報導").bold = True
        ranked = comp_stats.get(ind)
        if ranked:
            stat.add_run(";公司聲量前段:" + "、".join(
                f"{n}({c} 篇)" for n, c in ranked[:5]))
        stat.add_run("。重點報導如下:")
        for a in arts:
            _cite_para(doc, a)

    # ---- 建議策略 ----
    doc.add_page_break()
    _heading(doc, f"{next(section_no)}、建議策略", 1)
    disclaimer = doc.add_paragraph(
        "以下為系統依據期間內報導之主題統計所產生之策略觀察,屬 AI 分析,僅供參考,"
        "不構成投資或經營建議;每項觀察附代表性報導出處。")
    disclaimer.runs[0].font.color.rgb = GRAY
    if not themes:
        doc.add_paragraph("期間內資料量不足以形成跨產業主題觀察 —— 資料不足。")
    else:
        for theme, count, advice, samples in themes:
            p = doc.add_paragraph(style="List Number")
            p.add_run(f"{theme}(本期相關報導 {count} 篇):").bold = True
            p.add_run(advice)
            for s in samples:
                sp = doc.add_paragraph(style="List Bullet 2")
                sp.add_run(s["title"])
                c = sp.add_run(" " + format_citation(s))
                c.font.size = Pt(9)
                c.font.color.rgb = GRAY

    # ---- 資料缺口 ----
    if fetch_status:
        gaps = [s for s in fetch_status if s["status"] != "ok"]
        if gaps:
            _heading(doc, "資料缺口說明", 1)
            doc.add_paragraph("下列來源於本次更新未能成功抓取,相關產業之涵蓋度可能不足:")
            for g in gaps:
                doc.add_paragraph(
                    f"{g['source']}:{g['status']}({g.get('message') or '無訊息'},"
                    f"時間 {g['run_at']})", style="List Bullet")

    # ---- 附錄:引用來源總表 ----
    doc.add_page_break()
    _heading(doc, "附錄:引用來源總表", 1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, t in enumerate(("來源媒體", "標題", "發布日期", "原文連結")):
        hdr[i].text = t
    cited = {a["url"]: a for a in articles if a.get("industry")}
    for a in sorted(cited.values(), key=lambda x: (x["source"], x.get("published_at") or "")):
        row = table.add_row().cells
        row[0].text = a["source"]
        row[1].text = a["title"][:60]
        row[2].text = (a.get("published_at") or "")[:10]
        row[3].text = a["url"]
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)

    # ---- 輸出 ----
    if out_path is None:
        fname = f"{title}_{start_date}_{end_date}_{now.strftime('%Y%m%d%H%M')}.docx"
        out_path = REPORT_DIR / fname
    doc.save(out_path)
    return str(out_path)
